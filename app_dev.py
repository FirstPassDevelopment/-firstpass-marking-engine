import streamlit as st
from docx import Document
import pandas as pd
import io
import re
from statistics import mean
from openai import OpenAI
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo
from concurrent.futures import ThreadPoolExecutor, as_completed  # ✅ ADDED
import requests
import os

PASSWORD = os.getenv("APP_PASSWORD", "fallback-password")

if "auth" not in st.session_state:
    st.session_state.auth = False

if not st.session_state.auth:
    with st.form("login_form"):
        pw = st.text_input("Enter password", type="password")
        submitted = st.form_submit_button("Login")

    if submitted:
        if pw == PASSWORD:
            st.session_state.auth = True
            st.rerun()
        else:
            st.error("Incorrect password")

    st.stop()
# -----------------------
# ENV
# -----------------------

api_key = os.getenv("OPENAI_API_KEY")

if not api_key or not api_key.startswith("sk-"):
    st.error("Invalid or missing API key")
    st.stop()

client = OpenAI(api_key=api_key.strip())

# -----------------------
# INIT
# -----------------------

st.set_page_config(page_title="FirstPass Marking Engine – Stable", layout="wide")
st.title("FirstPass Marking Engine – Stable")
st.caption("Version 1.0 – Stable")

st.info("FirstPass provides an AI-assisted first-pass judgement. Final grading decisions remain the responsibility of the teacher.")

st.caption("Designed for Australian secondary classrooms using rubric-based marking aligned to ACARA standards.")

if "results" not in st.session_state:
    st.session_state.results = None

if "counts" not in st.session_state:
    st.session_state.counts = None

if "comments" not in st.session_state:
    st.session_state.comments = {}

if "comment_saved" not in st.session_state:
    st.session_state.comment_saved = {}

# -----------------------
# ACARA
# -----------------------

ACARA_EXPECTATIONS = {
    "Year 7": "Students explain how language features create meaning.",
    "Year 8": "Students explain how language features shape meaning.",
    "Year 9": "Students analyse how language features influence meaning.",
    "Year 10": "Students evaluate how language features shape meaning."
}

# -----------------------
# EXEMPLARS
# -----------------------

EXEMPLARS = """
C+ (At Standard):
- Clear and understandable ideas
- Basic explanation of language features
- Some structure
- Limited depth

IMPORTANT: Most students fall here.

---

B (Above Standard):
- Clearly more developed than typical work
- Consistent explanation of HOW meaning is created
- Deliberate use of evidence
- Controlled structure

Only award B if it is CLEARLY above standard.

---

A (Exceptional):
- Rare
- Insightful, precise, highly controlled
- Sophisticated reasoning sustained throughout

If unsure, it is NOT an A.
"""

# -----------------------
# HELPERS
# -----------------------

def extract_text(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def parse_feedback(feedback):
    sections = {"summary": "", "justification": "", "strengths": "", "improvements": ""}

    try:
        if "Summary:" in feedback:
            sections["summary"] = feedback.split("Summary:")[1].split("Justification:")[0].strip()
        if "Justification:" in feedback:
            sections["justification"] = feedback.split("Justification:")[1].split("Strengths:")[0].strip()
        if "Strengths:" in feedback:
            sections["strengths"] = feedback.split("Strengths:")[1].split("Areas for Improvement:")[0].strip()
        if "Areas for Improvement:" in feedback:
            sections["improvements"] = feedback.split("Areas for Improvement:")[1].strip()
    except:
        pass

    return sections

def send_feedback(feedback_text, results_data):
    endpoint = "https://formspree.io/f/maqardlg"

    payload = {
        "feedback": feedback_text,
        "results": str(results_data)
    }

    try:
        requests.post(endpoint, data=payload)
    except:
        pass

# ✅ HARD AU SPELLING ENFORCER
def enforce_au_spelling(text):
    replacements = {
        "analyze": "analyse",
        "analyzes": "analyses",
        "analyzed": "analysed",
        "analyzing": "analysing",
        "organization": "organisation",
        "organizations": "organisations",
        "color": "colour",
        "colors": "colours",
        "behavior": "behaviour",
        "center": "centre",
        "modeling": "modelling",
        "legalization": "legalisation",
        "legalize": "legalise",
        "legalized": "legalised",
        "legalizing": "legalising"
    }

    for us, au in replacements.items():
        text = re.sub(rf"\b{us}\b", au, text, flags=re.IGNORECASE)

    return text

# -----------------------
# PASS 1
# -----------------------

def summarise(text, task_type):

    prompt = f"""
Task Type: {task_type}

Summarise this response for marking.

Focus on:
- ideas
- language use
- structure
- depth

3–4 lines only.

{text}
"""

    res = client.responses.create(
    model="gpt-4.1-mini",
    input=prompt
)

    return res.output[0].content[0].text

# -----------------------
# PASS 2
# -----------------------

def grade_with_context(summary, cohort_summaries, year, task_type):

    prompt = f"""
You are marking a {year} English class.

LANGUAGE RULE:
Use Australian English conventions at all times.
- Use British/Australian spelling (e.g. analyse, organisation, colour)
- Avoid American spelling (e.g. analyze, organization, color)
- Maintain a professional Australian secondary teacher tone

Task Type: {task_type}

Curriculum expectation:
{ACARA_EXPECTATIONS[year]}

REFERENCE STANDARD:
{EXEMPLARS}

---

CRITICAL GRADING RULE:

Start at C+.

Then ask:

"Where does this realistically sit compared to typical student work?"

Move up based on quality:

C+ → B- → B → B+ → A- → A

IMPORTANT:
- You MAY move multiple bands if clearly justified.
- Do NOT artificially restrict movement.

If unsure:
- lean DOWN, not up

---

PERSUASIVE TASK ADJUSTMENT:

For persuasive writing, prioritise:

- strength of argument  
- clarity of position  
- effectiveness of persuasion  
- rhetorical appeals  
- audience impact  

Strong persuasion CAN justify B+ or A-

NEW RULE (A-BAND SENSITIVITY):

If the response is:
- consistently persuasive throughout
- clearly controlled and deliberate
- demonstrates strong audience impact

→ it SHOULD be considered for A-.

Do NOT cap clearly high-performing persuasive responses at B+.

---

DO NOT:
- reward fluency alone
- reward effort

FOCUS ON:
- effectiveness
- clarity
- control

---

Cohort:
{cohort_summaries}

---

Student:
{summary}

---

Output EXACTLY:

Overall Grade: [A+/A/A-/B+/B/B-/C+/C/C-/D/E]

Summary: [1 sentence]

Justification:
[2–3 sentences]

Strengths:
- point
- point

Areas for Improvement:
- point
- point
"""

    res = client.responses.create(
    model="gpt-4.1-mini",
    input=prompt
)

    return enforce_au_spelling(res.output[0].content[0].text)

# -----------------------
# INSIGHTS
# -----------------------

def generate_cohort_insights(results):

    all_feedback = "\n\n".join([r["feedback"] for r in results])

    prompt = f"""
You are an experienced Australian secondary English teacher.

LANGUAGE RULE:
Use Australian English conventions at all times.

IMPORTANT:
Do NOT contradict yourself.
If most responses show a strength, do NOT claim the same area as a major weakness.
Instead, qualify differences clearly (e.g. "while most responses showed X, several struggled with Y").

You have just marked THIS specific cohort.

You MUST:
- Identify patterns that are clearly visible in THESE responses
- Refer to actual trends (not generalisations)
- Avoid generic teacher-report language

STRICT RULES:

- Use phrases like:
  "Most responses..."
  "Several students..."
  "A small number of responses..."

- Every point MUST include:
  (1) the pattern  
  (2) how you know (what you saw repeatedly)

- DO NOT write generic statements like:
  "students demonstrated good understanding"
  "many students used evidence well"

- Be specific, grounded, and teacher-authentic

---

CRITICAL:

Instructional Priorities MUST include:
- what to reteach
- HOW to reteach it (specific strategy)
- what students will produce as evidence

---

FEEDBACK:
{all_feedback}

---

Output EXACTLY:

Class Strengths:
- pattern + how you know
- pattern + how you know

Class Weaknesses:
- pattern + how you know
- pattern + how you know

Instructional Priorities:
- reteach focus + how to teach it + what students will produce
- reteach focus + how to teach it + what students will produce
- reteach focus + how to teach it + what students will produce
"""

    res = client.responses.create(
    model="gpt-4.1-mini",
    input=prompt
)

    return enforce_au_spelling(res.output[0].content[0].text)

# -----------------------
# UI
# -----------------------

year = st.selectbox("Year", ["Year 7", "Year 8", "Year 9", "Year 10"])
task_type = st.selectbox("Task Type", ["Analytical","Creative","Hybrid","Persuasive","Short Response"])
task_file = st.file_uploader("Task Sheet (.docx)", type=["docx"])
rubric_file = st.file_uploader("Rubric (.docx)", type=["docx"])
student_files = st.file_uploader("Student Files", type=["docx"], accept_multiple_files=True)

status = st.empty()

if st.button("Run Marking"):

    summaries = []
    word_counts = []

    status.info("Reading responses...")

    for f in student_files:
        text = extract_text(f)
        wc = len(text.split())
        word_counts.append(wc)

        s = summarise(text, task_type)
        summaries.append((f.name, s, wc))

    avg_wc = mean(word_counts) if word_counts else 0

    status.info("Understanding cohort...")
    all_summaries = "\n\n".join([s[1] for s in summaries])

    results = []
    counts = {}

    status.info("Assigning grades (parallel processing)...")

    def process_student(name, summary, wc):
        feedback = grade_with_context(summary, all_summaries, year, task_type)

        grade = "C+"
        match = re.search(r"Overall Grade:\s*(A\+|A-|A|B\+|B-|B|C\+|C-|C|D|E)", feedback)
        if match:
            grade = match.group(1)

        length_flag = wc < (0.6 * avg_wc) if avg_wc > 0 else False

        return {
            "file": name,
            "grade": grade,
            "feedback": feedback,
            "word_count": wc,
            "length_flag": length_flag
        }

    progress_bar = st.progress(0)
    total = len(summaries)
    completed = 0

    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = [executor.submit(process_student, name, s, wc) for name, s, wc in summaries]

        for future in as_completed(futures):
            r = future.result()

            counts[r["grade"]] = counts.get(r["grade"], 0) + 1
            results.append(r)

            completed += 1
            progress_bar.progress(completed / total)

    results = sorted(results, key=lambda x: x["file"])

    status.success("Marking complete — results ready below.")

    st.session_state.results = results
    st.session_state.counts = counts

# -----------------------
# DISPLAY
# -----------------------

if st.session_state.results:

    st.write("### Distribution")
    st.write(st.session_state.counts)

    st.write("### Cohort Insights")
    st.write(generate_cohort_insights(st.session_state.results))

    for i, r in enumerate(st.session_state.results):
        st.write(f"Student {i+1} | {r['grade']}")

        if r["length_flag"]:
            st.warning("⚠️ This response is significantly shorter than the cohort")

        with st.expander("Feedback"):
            st.text(r["feedback"])

            key = f"comment_{i}"
            saved_key = f"saved_{i}"

            existing = st.session_state.comments.get(key, "")

            comment = st.text_input(
                "✏️ Teacher Comment (press Enter to save)",
                value=existing,
                key=key
            )

            if comment != existing:
                st.session_state.comments[key] = comment
                st.session_state.comment_saved[saved_key] = True

            if st.session_state.comment_saved.get(saved_key):
                st.success("Comment saved")
            else:
                st.warning("Not saved yet")

    st.write("### Feedback")

    feedback_input = st.text_area("Provide feedback on marking quality")

    if st.button("Submit Feedback"):
        send_feedback(feedback_input, st.session_state.results)
        st.success("Feedback sent")

# -----------------------
# EXCEL EXPORT
# -----------------------

if st.session_state.results:

    export_data = []

    for i, r in enumerate(st.session_state.results):
        parsed = parse_feedback(r["feedback"])

        export_data.append({
            "File": r["file"],
            "Grade": r["grade"],
            "Word Count": r["word_count"],
            "Length Flag": "Yes" if r["length_flag"] else "",
            "Summary": parsed["summary"],
            "Justification": parsed["justification"],
            "Strengths": parsed["strengths"],
            "Improvements": parsed["improvements"],
            "Teacher Comment": st.session_state.comments.get(f"comment_{i}", "")
        })

    df = pd.DataFrame(export_data)
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name="Results")
        sheet = writer.sheets["Results"]

        colour_map = {
            "A+": "C6EFCE","A": "E2F0D9","A-": "F3FBF3",
            "B+": "BDD7EE","B": "DDEBF7","B-": "F2F7FC",
            "C+": "FFE699","C": "FFF2CC","C-": "FFF9E6",
            "D": "F8CBAD","E": "F4CCCC"
        }

        for cell in sheet[1]:
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal="center")
            cell.fill = PatternFill(start_color="F5F5F5", fill_type="solid")

        widths = [30, 10, 14, 14, 40, 50, 40, 40, 45]
        for i, w in enumerate(widths):
            sheet.column_dimensions[chr(65+i)].width = w

        for row in sheet.iter_rows(min_row=2):
            grade_val = row[1].value
            fill = PatternFill(start_color=colour_map.get(grade_val, "FFFFFF"), fill_type="solid")

            for i, cell in enumerate(row):
                if i == 1:
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal="center", vertical="top")
                else:
                    cell.alignment = Alignment(wrap_text=True, vertical="top")

                cell.fill = fill

        last_row = len(df) + 1
        last_col = chr(65 + len(df.columns) - 1)

        table = Table(displayName="ResultsTable", ref=f"A1:{last_col}{last_row}")
        style = TableStyleInfo(name="TableStyleMedium9", showRowStripes=True)
        table.tableStyleInfo = style
        sheet.add_table(table)

    st.download_button("Download Excel", output.getvalue(), "firstpass_results.xlsx")